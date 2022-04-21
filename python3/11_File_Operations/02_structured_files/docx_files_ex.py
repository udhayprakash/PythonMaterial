#!/usr/bin/python
"""
Purpose: creating DOCX files
    pip install python-docx
"""
from docx import Document

document = Document()

# Adding a paragraph
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')

# It’s also possible to use one paragraph as a “cursor” and insert a new paragraph directly above it:
prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')

# Adding a heading
document.add_heading('The REAL meaning of the universe')

# sub-heading from level 1 to 9
document.add_heading('The role of dolphins', level=2)

# Adding a page break
document.add_page_break()

# Adding a table
table = document.add_table(rows=2, cols=2)

cell = table.cell(0, 1)
cell.text = 'parrot, possibly dead'

row = table.rows[1]
row.cells[0].text = 'Foo bar to you.'
row.cells[1].text = 'And a hearty foo bar to you too sir!'


for row in table.rows:
    for cell in row.cells:
        print(cell.text)

document.save('result.docx')
# Ref: https://python-docx.readthedocs.io/en/latest/user/quickstart.html

"""
Packages
---------
word Documentation - python-docx
Powerpoint Presentation - python-pptx
Excel/Spreadsheet       - openpyxl
PDF                     - Reportlab, python-pdfkit
"""
